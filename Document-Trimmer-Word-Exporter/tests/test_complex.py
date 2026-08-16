import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "runtime", "tests")
doc = Document()

# Section 1: pages 1-4, header SECTION A
for pageno in range(1, 5):
    t = doc.add_paragraph()
    r = t.add_run(f"PAGE {pageno} - SECTION A CONTENT")
    r.bold = True
    r.font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(1, 7):
        doc.add_paragraph(f"Section A page {pageno} paragraph {i}. ")
    if pageno < 4:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_paragraph()  # trailing empty paragraph on last page of section A

# Section 2: pages 5-6, header SECTION B
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
sec2.header.is_linked_to_previous = False
sec2.header.paragraphs[0].text = "SECTION B - CONFIDENTIAL"
for pageno in range(5, 7):
    t = doc.add_paragraph()
    t.add_run(f"PAGE {pageno} - SECTION B CONTENT").bold = True
    for i in range(1, 7):
        doc.add_paragraph(f"Section B page {pageno} paragraph {i}. ")
    if pageno < 6:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_paragraph()
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)  # trailing page break

# Section 3: pages 7-10, header SECTION C
sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
sec3.header.is_linked_to_previous = False
sec3.header.paragraphs[0].text = "SECTION C - TOP SECRET"
for pageno in range(7, 11):
    t = doc.add_paragraph()
    t.add_run(f"PAGE {pageno} - SECTION C CONTENT").bold = True
    for i in range(1, 7):
        doc.add_paragraph(f"Section C page {pageno} paragraph {i}. ")
    if pageno < 10:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_paragraph()
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        doc.add_paragraph()

src = os.path.join(OUT, "complex.docx")
doc.save(src)
print("saved complex.docx", os.path.getsize(src))

from engine import paginate, docx_trim
pc, bd = paginate.paginate_docx(src)
print("paginate: page_count =", pc)
print("boundaries:", bd)

for s, e in [(1, 4), (5, 6), (7, 10)]:
    out = os.path.join(OUT, f"complex_{s}_{e}.docx")
    docx_trim.split_docx_range(src, s, e, bd, out)
    print("split", s, e, "->", os.path.basename(out))
