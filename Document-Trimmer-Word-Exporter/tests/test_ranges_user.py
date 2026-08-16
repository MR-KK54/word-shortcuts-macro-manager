import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document
from docx.enum.text import WD_BREAK
from engine import paginate, docx_trim

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "runtime", "tests")
src = os.path.join(OUT, "user_range_5p.docx")

doc = Document()
for i in range(1, 6):
    doc.add_paragraph(f"Page {i} paragraph A - heading content")
    doc.add_paragraph(f"Page {i} paragraph B - body text content")
    if i < 5:
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)  # explicit page break

doc.save(src)
print("Saved 5-page test docx:", src)

pc, bd = paginate.paginate_docx(src)
print("Source page count:", pc)
print("Boundaries:", bd)

success = True
for end in (1, 2, 3, 4, 5):
    out = os.path.join(OUT, f"user_split_1_{end}.docx")
    docx_trim.split_docx_range(src, 1, end, bd, out)
    out_pc, _ = paginate.paginate_docx(out)
    match = "PASS" if out_pc == end else "FAIL"
    if out_pc != end:
        success = False
    print(f"Range 1-{end} -> Expected {end} page(s), got {out_pc} page(s). [{match}]")

if not success:
    sys.exit(1)
