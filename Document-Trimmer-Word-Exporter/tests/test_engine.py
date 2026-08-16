import os, sys, shutil, zipfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, Inches

OUT = os.path.join(os.path.dirname(__file__), "runtime", "tests")
shutil.rmtree(OUT, ignore_errors=True)
os.makedirs(OUT, exist_ok=True)

doc = Document()

# --- Header / footer ---
sec = doc.sections[0]
header = sec.header
hp = header.paragraphs[0]
hp.text = "CONFIDENTIAL - PROJECT REPORT"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer = sec.footer
fp = footer.paragraphs[0]
fp.text = "Page Footer - Internal Use Only"
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for pageno in range(1, 5):
    title = doc.add_paragraph()
    run = title.add_run(f"PAGE {pageno} - SECTION TITLE")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(1, 9):
        p = doc.add_paragraph()
        r = p.add_run(f"Page {pageno} body paragraph {i}. ")
        r.italic = (i % 2 == 0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)

    # A small table per page
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    for ri in range(3):
        for ci in range(3):
            tbl.rows[ri].cells[ci].text = f"P{pageno}R{ri}C{ci}"

    # trailing junk on the last page: empty paragraphs + explicit page breaks
    if pageno in (2, 3):
        doc.add_paragraph()  # empty paragraph
        b = doc.add_paragraph()
        b.add_run().add_break(WD_BREAK.PAGE)  # page break
    else:
        doc.add_paragraph()  # trailing empty paragraph at end of doc

src = os.path.join(OUT, "sample.docx")
doc.save(src)
print("saved", src, os.path.getsize(src), "bytes")

# Now split/paginate using the engine
from engine import paginate, docx_trim

page_count, boundaries = paginate.paginate_docx(src)
print("PAGINATED: page_count =", page_count)
print("boundaries (page -> last unit idx):", boundaries)

for start, end in [(1, 2), (3, 3), (2, 3)]:
    out = os.path.join(OUT, f"split_{start}_{end}.docx")
    docx_trim.split_docx_range(src, start, end, boundaries, out)
    print("split", start, end, "->", os.path.basename(out), os.path.getsize(out), "bytes")
