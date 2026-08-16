import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import docx
from docx.shared import Pt
from engine import paginate, docx_trim, verify

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "runtime", "tests")
os.makedirs(OUT, exist_ok=True)

path = os.path.join(OUT, "verify_loop_src.docx")
d = docx.Document()
d.add_heading("Page 1 header", level=1)
for para in range(1, 40):
    p = d.add_paragraph()
    run = p.add_run(f"Paragraph {para}. ")
    run.font.size = Pt(12)
    for w in range(1, 30):
        run.add_text(f"word{w} ")
d.add_page_break()
d.add_heading("Second section", level=1)
for para in range(40, 70):
    p = d.add_paragraph()
    run = p.add_run(f"Body {para}. ")
    run.font.size = Pt(12)
    for w in range(1, 30):
        run.add_text(f"tok{w} ")
d.save(path)

page_count, boundaries = paginate.paginate_docx(path)
print(f"source: {page_count} pages, boundaries={boundaries}")

total = 0
fail = 0
start_pages = [1] + [p for p in range(2, page_count + 1)]
for start in start_pages[:3]:
    for end in range(start, min(start + 2, page_count + 1)):
        total += 1
        out = os.path.join(OUT, f"verify_loop_{start}_{end}.docx")
        docx_trim.split_docx_range(path, start, end, boundaries, out)
        report = verify.verify_export(path, out, start, end)
        status = "PASS" if report["pass"] else "FAIL"
        if not report["pass"]:
            fail += 1
        print(f"  pages {start}-{end}: {status} "
              f"(expected {report['page_count']['expected']}, "
              f"got {report['page_count']['actual']}, corrected {report['corrected']})")
        for e in report["pages"]:
            if not e["match"] or not e["text_match"]:
                print(f"      p{e['page']}: match={e['match']} text={e['text_match']} {e['detail']}")

print(f"\n{total - fail}/{total} splits passed verification")
if fail:
    print("RESULT: FAILURES DETECTED")
    sys.exit(1)
print("RESULT: ALL PASS")
