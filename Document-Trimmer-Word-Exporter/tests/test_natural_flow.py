import os, sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import docx
from docx.shared import Pt
from engine import word_com, word_markers, paginate, docx_trim

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "runtime", "tests")
path = os.path.join(OUT, "natural_flow.docx")

d = docx.Document()
for para in range(1, 60):
    p = d.add_paragraph()
    run = p.add_run(f"Paragraph {para}. ")
    run.font.size = Pt(12)
    for w in range(1, 30):
        run.add_text(f"word{w} ")
d.save(path)

# Word COM pagination (ground truth)
word = word_com._open_word()
try:
    doc = word.Documents.Open(FileName=os.path.abspath(path), ReadOnly=True, AddToRecentFiles=False, Visible=False)
    doc.Repaginate()
    word_pages = int(doc.ComputeStatistics(2))
    doc.Close(SaveChanges=0)
finally:
    word.Quit()

# Marker-based (what Render would use)
m = word_markers.paginate(path)
e = paginate._explicit_break_pagination(path)

# What the engine produces: split pages 1-3, count with Word
from engine import docx_trim as dt
OUT2 = os.path.join(OUT, "natural_split.docx")
try:
    dt.split_docx_range(path, 1, 3, m[1] if m else e[1], OUT2)
    word = word_com._open_word()
    try:
        doc = word.Documents.Open(FileName=os.path.abspath(OUT2), ReadOnly=True, AddToRecentFiles=False, Visible=False)
        doc.Repaginate()
        out_pages = int(doc.ComputeStatistics(2))
        doc.Close(SaveChanges=0)
    finally:
        word.Quit()
except Exception as ex:
    out_pages = f"ERR {ex}"

print(f"natural_flow.docx: Word pages={word_pages}")
print(f"  markers pagination: {m[0] if m else None} pages, boundaries={m[1] if m else None}")
print(f"  explicit pagination: {e[0]} pages")
print(f"  split 1-3 -> output opens with {out_pages} page(s) in Word")
